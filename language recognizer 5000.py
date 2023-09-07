import os
import re
from docx import Document
from unidecode import unidecode
from collections import Counter
import json


def find_path(file, directory):

	"""
	returns path to given file in your pc, where 'file' is the name of the file and 'directory' is the path to the
	directory where it is located
	"""

	file_path = os.path.join(directory, file)
	return os.path.abspath(file_path)


def remove_diacritics(text):

	"""
	removes diacritics from text using unidecode module
	"""

	cleaned_text = unidecode(text)
	return cleaned_text


def remove_whitespace(text):

	"""
	removes whitespace from text
	"""

	cleaned_text = ''.join(text.split())
	return cleaned_text


def remove_non_letter_characters(text):

	"""
	replaces all non-letter characters with empty strings
	(for specifics check regex syntax)
	"""

	cleaned_text = re.sub(r'[^a-zA-Z]', '', text)

	return cleaned_text


def clean_text(text):

	"""
	returns text as lowercase, without diacritics, whitespaces or non-letter characters
	"""

	text = remove_diacritics(text)
	text = remove_whitespace(text)
	text = remove_non_letter_characters(text)
	text = text.lower()
	return text


def docx_pure_text(file):

	"""
	extracts pure text from docx file
	(lowercase, without diacritics, whitespaces or non-letter characters)
	"""

	doc = Document(file)
	full_text = []

	for paragraph in doc.paragraphs:
		text = clean_text(paragraph.text)
		full_text.append(text)

	return ''.join(full_text)


def create_docx(text, path):

	"""
	creates a docx from given text in a directory given by the path
	"""

	doc = Document()
	doc.add_paragraph(text)
	doc.save(path)


def extract_and_rewrite(file, directory):

	"""
	extracts pure text from docx file and rewrites the file
	"""

	path = find_path(file, directory)
	text = docx_pure_text(path)
	create_docx(text, path)  # rewrites the original docx as pure text


def rewrite_files(path):

	"""
	rewrites all docx files in directory (whose path is given) to contain only pure text
	"""

	for file in os.listdir(path):
		if os.path.isfile(os.path.join(path, file)) and os.path.splitext(file)[1].lower() == '.docx':
			extract_and_rewrite(file, path)


def descendsort_dict(dictionary: dict):

	"""
	sorts dictionary by values in descending order
	"""

	return dict(sorted(dictionary.items(), key=lambda item: item[1], reverse=True))


class DatasetBuilder:

	def __init__(self):
		pass  # technicality to avoid IndentationError

	def create_dataset(self, languages, directory):

		"""
		creates a dataset to be used to build language profiles,
		'languages' is a dictionary of this form:
		{
		'EN': {'name': 'English'},
		'FR': {'name': 'French'},
		},
		'directory' is the path to the directory from which we want to read our data,
		dummy dataset to illustrate the desired structure:
		{
		CZ: {'name': 'Czech', 'texts': [ 'hody hody doprovody', 'dejte vejce malovany' ]}
		FR: {'name': 'French', 'texts': [ 'ta grandmere en string', 'me met l eau a la bouche' ]}
		}
		"""

		dataset = {}
		for l_code, l_data in languages.items():
			l_name = l_data['name']
			l_texts = []
			l_directory = os.path.join(directory, l_code)
			if os.path.exists(l_directory) and os.path.isdir(l_directory):
				l_texts = self.read_docx_files_in_directory(l_directory)
			dataset[l_code] = {'name': l_name, 'texts': l_texts}
		return dataset

	def read_docx_files_in_directory(self, directory):

		"""
		reads docx files in given directory and copies their content to a python list
		"""

		texts = []
		for file in os.listdir(directory):
			if file.endswith(".docx"):
				file_path = os.path.join(directory, file)
				doc = Document(file_path)
				text = "".join([paragraph.text for paragraph in doc.paragraphs])
				texts.append(text)
		return texts


class LanguageDetector:
	# in this class, 'l' or 'lang' is often used as short for 'language'
	def __init__(self, languages=None):
		self.l_profiles = languages if languages else {}  # either dictionary containing language profile(s) or empty dict
		self.n = 3

	def add_l(self, l_code, l_name, ngrams):

		"""
		adds language profile,
		for example for spanish l_code is 'ES' and l_name is 'Spanish'
		"""

		self.l_profiles[l_code] = {'name': f'{l_name}', 'ngrams': ngrams}

	def build_l_profiles(self, dataset: dict):

		"""
		builds up language profiles (includes the code, name, ngrams and their frequencies) from given data
		"""

		for l_code, l_data in dataset.items():
			texts = l_data['texts']
			ngrams = self.extract_ngrams(texts)
			total_ngrams = sum(ngrams.values())
			ngrams_prob = {ngram: freq / total_ngrams for ngram, freq in ngrams.items()}  # dictionary containing the
			# probability of occurrence of each ngram in our dataset
			self.add_l(l_code, l_data['name'], ngrams_prob)

	def extract_ngrams(self, text):

		"""
		extracts ngrams from given text
		"""

		ngrams = Counter()
		for i in range(len(text) - self.n + 1):
			ngram = text[i:i + self.n]
			ngrams[ngram] += 1
		return ngrams

	def calculate_score(self, text_ngrams, l_ngrams):

		"""
		calculates score for each language given the ngrams of the input text and language,
		further explanation of scoring system given in user documentation
		"""

		score = 0
		for ngram, freq in text_ngrams.items():
			if ngram in l_ngrams:
				score += freq * (l_ngrams[ngram])
		return score

	def detect_language(self, text):

		"""
		detects the language the input text is written in
		"""

		text = clean_text(text)
		length = len(text)
		if length < 30:
			print(
				'minimum text length is 30 letters (non-letter characters such as numbers, punctuation or whitespaces do not count)')

		else:
			text_ngrams = self.extract_ngrams(text)
			l_scores = {l_code: self.calculate_score(text_ngrams, l_data['ngrams']) for l_code, l_data in self.l_profiles.items()}
			lang = max(l_scores, key=l_scores.get)  # code of the language with the highest score
			lang_score = l_scores[lang]  # score of the language with the highest score
			if (30 <= length <= 50 and lang_score < 0.015) or (
					50 < length <= 200 and lang_score < 0.03) or (length > 200 and lang_score < 0.13):
				print('language not detected')
				print()
				print(descendsort_dict(l_scores))
			else:
				return self.l_profiles[lang]['name'], descendsort_dict(l_scores)  # returns name of detected language & the dictionary of scores of all the languages

	def get_results(self, text):

		"""
		prints input text, detected language and the dictionary of language scores for all the languages
		"""

		lang = self.detect_language(text)
		print()
		print('input text:', text)
		if lang:
			print()
			print('detected language:', lang[0])
			print()
			print('all language scores:', lang[1])


class JSONLanguageProfiles:

	def __init__(self):
		script_directory = os.path.dirname(os.path.abspath(__file__))  # directory where the python script and file to
		# store JSON files is located
		self.profiles_path = os.path.join(script_directory, 'JSON_language_profiles')  # path to directory with JSON files

	def get_profile_filename(self, lang_code):

		"""
		returns path to given language profiles JSON file
		"""

		return os.path.join(self.profiles_path, f"{lang_code}_profile.json")

	def save_profile(self, l_code, data):

		"""
		saves language profile to JSON file
		"""

		file_name = f'{l_code}_profile.json'
		with open(file_name, 'w') as file:
			json.dump(data, file, indent=4)  # indentation for user readability, equivalent of nesting w 4 spaces in a script

	def load_profile(self, l_code):

		"""
		parses given language profile's JSON file into a python dictionary
		"""

		file_name = self.get_profile_filename(l_code)
		try:
			with open(file_name, 'r') as file:
				return json.load(file)
		except FileNotFoundError:
			return None


def main():
	languages = {
		'EN': {'name': 'English'},
		'FR': {'name': 'French'},
		'CZ': {'name': 'Czech'},
		'PL': {'name': 'Polish'},
		'LT': {'name': 'Lithuanian'},
		'NL': {'name': 'Dutch'},
		'DE': {'name': 'German'},
		'ES': {'name': 'Spanish'},
		'IT': {'name': 'Italian'},
		'HU': {'name': 'Hungarian'}
	}

	ld = LanguageDetector()
	language_profiles = JSONLanguageProfiles()

	for l_code in languages.keys():
		profile = language_profiles.load_profile(l_code)
		if profile:
			ld.add_l(l_code, profile['name'], profile['ngrams'])

	text = str(input('Enter the text to be detected:'))
	ld.get_results(text)


if __name__ == '__main__':
	main()
